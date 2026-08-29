"""SAE / IEC report generator — produces the PGIMER-DSMC Serious Adverse
Event Reporting Form (CDSCO / NDCT Rules 2019, 22 items) and the PI
covering letter as .docx files, pre-filled from a SAEReport row plus the
context the caller assembles from the clinical database.

Trial- and site-level constants come from sae_config. Anything not known
is printed as an explicit "[TO BE PROVIDED — ...]" blank, never guessed.
The randomised oxygen arm is never revealed (see sae_config TRIAL
suspect_intervention).
"""

from __future__ import annotations

import io
from datetime import date, datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sae_config as cfg

_NA = "Not applicable"
_BLANK = "—"


# --------------------------------------------------------------------------
# small helpers
# --------------------------------------------------------------------------

def _v(x):
    if x is None:
        return _BLANK
    s = str(x).strip()
    return s or _BLANK


def _fmt_dt(x):
    if not x:
        return _BLANK
    s = str(x)
    s = s.replace("T", " ")
    return s


def _yesno_list(x):
    if isinstance(x, (list, tuple)):
        return ", ".join(str(i) for i in x if i) or _BLANK
    return _v(x)


def _doc():
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    for section in d.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
    return d


def _h(d, text, size=13, space_before=10):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def _para(d, text, bold=False, italic=False, size=10.5):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def _kv_table(d, rows):
    """Two-column label/value table."""
    t = d.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    for label, value in rows:
        cells = t.add_row().cells
        cells[0].width = Inches(2.7)
        cells[1].width = Inches(4.0)
        rp = cells[0].paragraphs[0].add_run(label)
        rp.bold = True
        rp.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(_v(value)).font.size = Pt(10)
    return t


def _long_answer(d, prompt, value):
    _para(d, prompt, italic=True, size=9.5)
    box = d.add_paragraph()
    box.paragraph_format.space_after = Pt(8)
    box.add_run(_v(value)).font.size = Pt(10.5)


# --------------------------------------------------------------------------
# context shape (assembled by the endpoint)
# --------------------------------------------------------------------------
# ctx = {
#   "site_code": "PGIMER",
#   "patient": {"initials", "identifier", "gender", "dob", "age_text",
#               "weight_kg", "gestation"},
#   "concomitant": [{"therapy", "first_day", "last_day"}, ...],
#   "prior_reports": [{"report_type", "report_date", "diary_no"}, ...],
#   "linked_ae": {"description", "grade_label", "evidence"} | None,
#   "generated_by": "user@site",
# }


def _patient_rows(patient):
    p = patient or {}
    return [
        ("Initials / identifier", p.get("initials") or _BLANK),
        ("Hospital / OPD record number", p.get("identifier") or _BLANK),
        ("Gender", p.get("gender") or _BLANK),
        ("Age / date of birth", p.get("age_text") or p.get("dob") or _BLANK),
        ("Weight", p.get("weight_kg") or _BLANK),
        ("Gestation at birth", p.get("gestation") or _BLANK),
    ]


def _suspect_rows():
    s = cfg.TRIAL["suspect_intervention"]
    return [
        ("Generic name of the intervention", s["generic_name"]),
        ("Indication", s["indication"]),
        ("Dosage form and strength", s["dosage_form"]),
        ("Daily dose and regimen", s["daily_dose"]),
        ("Route of administration", s["route"]),
        ("Start date/time", s["start"]),
        ("Stop date/time or duration", s["stop"]),
    ]


def _concomitant_table(d, concomitant):
    if not concomitant:
        _para(d, "None recorded in the daily study logs for this enrolment, or "
                 "the logs are incomplete — the investigator should complete "
                 "this section.", italic=True, size=9.5)
        return
    t = d.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, label in enumerate(("Therapy / drug", "First day given", "Last day given", "Notes")):
        r = hdr[i].paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in concomitant:
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(_v(row.get("therapy"))).font.size = Pt(9.5)
        cells[1].paragraphs[0].add_run(_v(row.get("first_day"))).font.size = Pt(9.5)
        cells[2].paragraphs[0].add_run(_v(row.get("last_day"))).font.size = Pt(9.5)
        cells[3].paragraphs[0].add_run(_v(row.get("notes"))).font.size = Pt(9.5)


# --------------------------------------------------------------------------
# the SAE reporting form
# --------------------------------------------------------------------------

def build_sae_report_docx(sae, ctx) -> bytes:
    ctx = ctx or {}
    site = cfg.site_for(ctx.get("site_code"))
    iec = cfg.iec_for(ctx.get("site_code"))
    d = _doc()

    title = d.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Serious Adverse Event Reporting Form")
    tr.bold = True
    tr.font.size = Pt(15)
    sub = d.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("PGIMER Institute Ethics Committee — CDSCO / NDCT Rules 2019 format")
    sr.italic = True
    sr.font.size = Pt(9.5)

    _para(d, cfg.SUBMISSION_RULES, italic=True, size=9)

    death = "Death" in (sae.seriousness or []) if isinstance(sae.seriousness, list) else False

    _h(d, "1–10. Report and trial identification")
    _kv_table(d, [
        ("1. Country where the SAE occurred", cfg.TRIAL["country"]),
        ("2. SAE report of", "Death" if death else "Other than death"),
        ("   If other than death, specify", _NA if death else _v(sae.diagnosis)),
        ("3. Protocol title", cfg.TRIAL["protocol_title"]),
        ("4. Protocol Study No / ID / Code", cfg.TRIAL["protocol_no"]),
        ("5. CDSCO clinical trial permission", cfg.TRIAL["cdsco_permission"]),
        ("6. CTRI registration number", cfg.TRIAL["ctri_no"]),
        ("7. Sponsor", f'{cfg.TRIAL["sponsor"]["name"]}\n{cfg.TRIAL["sponsor"]["address"]}\n'
                       f'{cfg.TRIAL["sponsor"]["contact"]} · {cfg.TRIAL["sponsor"]["email"]}'),
        ("8. CRO", cfg.TRIAL["cro"]),
        ("9. Initial / Follow-up", _v(sae.report_type)),
        ("10. If follow-up: date & diary no. of initial / recent report",
         _prior_ref(ctx.get("prior_reports"))),
    ])

    _h(d, "11–12. Patient details")
    _kv_table(d, _patient_rows(ctx.get("patient")))
    _para(d, f"Study enrolment ID: {_v(sae.enrollment_id)}", size=9.5)

    _h(d, "13. Suspected intervention")
    _kv_table(d, _suspect_rows())

    _h(d, "14. Other treatment(s) — concomitant drugs and non-drug therapies")
    _concomitant_table(d, ctx.get("concomitant"))

    _h(d, "15. Details of the suspected adverse reaction")
    la = ctx.get("linked_ae")
    desc = sae.narrative or ""
    if la:
        ev = f"\n\nLinked recorded adverse event: {la.get('description')}" \
             f" (auto-detected severity {la.get('grade_label')})." \
             f"\nDetection evidence: {la.get('evidence')}"
        desc = (desc + ev).strip()
    _long_answer(d,
        "15.1  Full description of the reaction including body site and severity, the "
        "criterion/criteria for regarding the report as serious, a specific diagnosis, "
        "and the chronology of events:",
        desc)
    _kv_table(d, [
        ("15.2  Start date/time of onset", _fmt_dt(sae.onset_datetime)),
        ("15.3  Stop date/time or duration",
         "Ongoing" if getattr(sae, "ongoing", False) else _fmt_dt(sae.end_datetime)),
        ("15.4  Dechallenge / rechallenge",
         "Not applicable — the randomised oxygen intervention is titrated to "
         "physiological targets, not withdrawn/re-administered as a discrete challenge; "
         "see the respiratory support log."),
        ("15.5  Setting", "Neonatal Intensive Care Unit / delivery room"),
        ("15.6  Seriousness criteria met", _yesno_list(sae.seriousness)),
        ("15.7  Severity (INC NAESS)", cfg.severity_label(sae.severity)),
        ("15.8  Action taken", _v(sae.action_taken)),
        ("15.9  If there was a delay in reporting, the reason",
         _v(getattr(sae, "reporting_delay_reason", None))),
    ])

    _h(d, "16. Outcome")
    _long_answer(d, "16.1  Recovery and any sequelae; results of specific tests and/or "
                    "treatment conducted:", _outcome_line(sae))
    _long_answer(d, "16.2  For a fatal outcome — cause of death, a comment on its possible "
                    "relationship to the suspected reaction, and any post-mortem findings:",
                 _v(getattr(sae, "death_details", None)) if death else _NA)
    _long_answer(d, "16.3  Other relevant information — medical history, allergy, drug/alcohol "
                    "exposure, family history, findings from special investigations:",
                 _v(getattr(sae, "other_relevant_info", None)))

    _h(d, "17. Investigator")
    _kv_table(d, [
        ("CT site number", site["ct_site_number"]),
        ("Name", _v(sae.investigator_name) if sae.investigator_name else site["pi_name"]),
        ("Address", site["address"]),
        ("Telephone / mobile (e-mail)", f'{site["pi_phone"]} ({site["pi_email"]})'),
        ("Profession (specialty)", site["pi_specialty"]),
        ("Date reported to the Licensing Authority (CDSCO)",
         _v(ctx.get("date_reported_cdsco"))),
        ("Date reported to the Ethics Committee overseeing the site",
         _v(sae.investigator_date or ctx.get("date_reported_ec"))),
        ("Signature of the investigator", "________________________"),
    ])

    _h(d, "18. Ethics committee")
    _kv_table(d, [
        ("a. Name and address", f'{iec["name"]}\n{iec["address"]}'),
        ("b. Name of Chairperson and address",
         f'{iec.get("chairperson", cfg.PH)}\n{iec.get("chairperson_address", cfg.PH)}'),
        ("c. Telephone / mobile", iec.get("phone", cfg.PH)),
        ("d. E-mail", ", ".join(iec.get("emails", []))),
    ])

    _h(d, "19–21. Causality and compensation")
    _kv_table(d, [
        ("19. Causality assessment by Investigator (Related / Unrelated)", _v(sae.causality)),
        ("20. Causality assessment by Sponsor / CRO (Related / Unrelated)",
         _v(getattr(sae, "sponsor_causality", None))),
        ("21. Details of compensation provided for injury or death "
         "(if none, the reason)", cfg.TRIAL["compensation_clause"]),
    ])

    _h(d, "22. Attachments")
    _para(d, "a) Laboratory investigation report / discharge summary (if available "
             "and applicable): _______________________________")
    _para(d, "b) Post-mortem report (if applicable) / any additional document: "
             "_______________________________")

    d.add_paragraph()
    _para(d, "Signature of the investigator: ________________________     Date: ____________")

    foot = d.add_paragraph()
    fr = foot.add_run(
        f"\nGenerated by the PORTAL Trial data system on "
        f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
        + (f" · {ctx.get('generated_by')}" if ctx.get("generated_by") else "")
        + ". Review every field, complete all [TO BE PROVIDED] items, and attach "
        "supporting documents before submission."
    )
    fr.italic = True
    fr.font.size = Pt(8)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _prior_ref(prior):
    if not prior:
        return _NA
    parts = []
    for r in prior:
        parts.append(
            f'{r.get("report_type", "?")} on {r.get("report_date", "?")}'
            + (f' (Diary no. {r["diary_no"]})' if r.get("diary_no") else "")
        )
    return "; ".join(parts)


def _outcome_line(sae):
    parts = []
    if sae.outcome:
        parts.append(f"Outcome: {sae.outcome}.")
    if getattr(sae, "outcome_details", None):
        parts.append(sae.outcome_details)
    if getattr(sae, "date_of_death", None):
        parts.append(f"Date of death: {sae.date_of_death}.")
    return " ".join(parts) if parts else _BLANK


# --------------------------------------------------------------------------
# the PI covering letter (24-hour initial notification)
# --------------------------------------------------------------------------

def build_covering_letter_docx(sae, ctx) -> bytes:
    ctx = ctx or {}
    site = cfg.site_for(ctx.get("site_code"))
    iec = cfg.iec_for(ctx.get("site_code"))
    d = _doc()

    today = date.today().isoformat()
    _para(d, f"Date: {today}")
    d.add_paragraph()
    _para(d, "To,")
    _para(d, iec.get("addressee", "Chairperson, Institute Ethics Committee"))
    _para(d, iec["name"])
    _para(d, iec["address"])
    d.add_paragraph()

    death = "Death" in (sae.seriousness or []) if isinstance(sae.seriousness, list) else False
    _para(d, "Subject: 24-hour notification of a Serious Adverse Event — "
             + cfg.TRIAL["protocol_title"], bold=True)
    d.add_paragraph()

    _para(d, "Respected Chairperson,")
    d.add_paragraph()
    _para(d,
        f"I write to notify the Institute Ethics Committee, within 24 hours as required, "
        f"of a Serious Adverse Event in the above trial "
        f"(Protocol No. {cfg.TRIAL['protocol_no']}; CTRI No. {cfg.TRIAL['ctri_no']}).")
    d.add_paragraph()

    _kv_table(d, [
        ("Site", f'{site["display"]} (CT site no. {site["ct_site_number"]})'),
        ("Study enrolment ID", _v(sae.enrollment_id)),
        ("Event term / diagnosis", _v(sae.diagnosis)),
        ("Category", "Death" if death else "Serious — other than death"),
        ("Seriousness criteria met", _yesno_list(sae.seriousness)),
        ("Date/time of onset", _fmt_dt(sae.onset_datetime)),
        ("Date the investigator became aware", _v(ctx.get("date_aware") or sae.report_date)),
        ("Severity (INC NAESS)", cfg.severity_label(sae.severity)),
        ("Current status / outcome",
         "Ongoing" if getattr(sae, "ongoing", False) else _v(sae.outcome)),
    ])
    d.add_paragraph()

    _para(d,
        "The study intervention is a blinded randomised delivery-room oxygen "
        "concentration; the site team and this notification remain blinded to the "
        "allocated arm. A detailed Serious Adverse Event Reporting Form will be "
        "submitted within 14 days of the event as required. The event has also been / "
        "will be reported to the Sponsor and to CDSCO within the applicable timeframes.")
    d.add_paragraph()

    _para(d, "Thank you.")
    d.add_paragraph()
    _para(d, "Yours sincerely,")
    d.add_paragraph()
    _para(d, _v(sae.investigator_name) if sae.investigator_name else site["pi_name"])
    _para(d, f'Principal Investigator, {site["display"]}')
    _para(d, f'{site["pi_phone"]} · {site["pi_email"]}')

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
